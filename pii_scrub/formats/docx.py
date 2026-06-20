"""Word .docx — scrub paragraph and table text, keep document structure.

Optional dependency: python-docx  (pip install 'pii-airlock[docx]')

Text is scrubbed at paragraph granularity so that entities spanning multiple
runs are still caught; when a paragraph changes, its runs are collapsed into a
single run (formatting within a scrubbed paragraph is not preserved).
"""

from __future__ import annotations

import io

from .base import Handler, MissingExtra, Scrubber


class DocxHandler(Handler):
    reads_binary = True
    binary_output = True
    out_suffix = ".docx"

    def scrub(self, scrub: Scrubber) -> bytes:
        try:
            import docx  # type: ignore
        except ImportError as exc:  # pragma: no cover - depends on install
            raise MissingExtra("docx", "python-docx") from exc

        assert isinstance(self.raw, (bytes, bytearray))
        doc = docx.Document(io.BytesIO(self.raw))

        for para in doc.paragraphs:
            _scrub_paragraph(para, scrub)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _scrub_paragraph(para, scrub)

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()


def _scrub_paragraph(para, scrub: Scrubber) -> None:
    original = para.text
    if not original.strip():
        return
    cleaned = scrub(original)
    if cleaned == original:
        return
    # Collapse into the first run; drop the rest.
    if para.runs:
        para.runs[0].text = cleaned
        for run in para.runs[1:]:
            run.text = ""
    else:  # pragma: no cover - paragraphs normally have runs
        para.add_run(cleaned)
